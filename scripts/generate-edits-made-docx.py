"""
Build Edits_made.docx from Edits_made.md (YAML frontmatter stripped; ## / ### / bullets).
Run from repo root: python scripts/generate-edits-made-docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path


def strip_frontmatter(text: str) -> str:
    lines = text.splitlines()
    if not lines or lines[0].strip() != "---":
        return text
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            return "\n".join(lines[i + 1 :]).lstrip("\n")
    return text


def add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "Edits_made.md"
    out_path = root / "Edits_made.docx"
    if not md_path.is_file():
        print(f"Missing {md_path}", file=sys.stderr)
        return 1

    try:
        from docx import Document
    except ImportError:
        print("Install: pip install python-docx", file=sys.stderr)
        return 1

    md = strip_frontmatter(md_path.read_text(encoding="utf-8"))
    doc = Document()
    doc.core_properties.title = "voiceVault — Edits made"
    doc.core_properties.comments = "Generated from Edits_made.md"

    for raw in md.splitlines():
        s = raw.rstrip()
        if not s:
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, s[2:])
        else:
            p = doc.add_paragraph()
            add_runs_with_bold(p, s)

    doc.save(out_path)
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
