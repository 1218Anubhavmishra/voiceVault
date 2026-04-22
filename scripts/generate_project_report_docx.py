from __future__ import annotations

import re
from pathlib import Path


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "project-report.md"
    out_path = root / "project-report.docx"
    assets_dir = root / "report-assets"

    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "Missing dependency python-docx. Install it with:\n"
            r".\.venv\Scripts\python.exe -m pip install python-docx"
        ) from e

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    in_code = False
    current_list: str | None = None  # "ul" only for now

    def end_list() -> None:
        nonlocal current_list
        current_list = None

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            in_code = not in_code
            end_list()
            continue

        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            continue

        if line.strip() == "":
            end_list()
            continue

        if line.startswith("---") and (doc.paragraphs and doc.paragraphs[-1].text == ""):
            # Ignore frontmatter separators if they appear as standalone.
            continue

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img:
            end_list()
            alt = (img.group(1) or "").strip()
            src = (img.group(2) or "").strip()
            src_path = (root / src).resolve() if not Path(src).is_absolute() else Path(src)
            if src_path.exists():
                if alt:
                    doc.add_paragraph(alt)
                # Keep images a reasonable width for typical pages.
                doc.add_picture(str(src_path), width=Inches(6.5))
            else:
                doc.add_paragraph(f"[Missing image: {src}]")
            continue

        m = re.match(r"^(#{2,6})\s+(.*)$", line)
        if m:
            end_list()
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 2:
                doc.add_heading(title, level=1)
            elif level == 3:
                doc.add_heading(title, level=2)
            else:
                doc.add_heading(title, level=3)
            continue

        if line.startswith("- "):
            current_list = current_list or "ul"
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        end_list()

        # Very small amount of inline markdown cleanup:
        # - `code` -> code
        # - **bold** -> bold
        cleaned = re.sub(r"`([^`]+)`", r"\1", line)
        cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
        doc.add_paragraph(cleaned)

    doc.save(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

